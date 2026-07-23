#!/usr/bin/env python3
"""Heuristic statistical reporting inventory and consistency flags for DOCX manuscripts."""
from __future__ import annotations

import argparse
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document

PATTERNS = {
    "sample_sizes": re.compile(r"\b(?:N|n)\s*=\s*([0-9][0-9,]*)"),
    "p_values": re.compile(r"\bp\s*(=|<|>|≤|≥)\s*(\.?\d+)", re.I),
    "confidence_intervals": re.compile(r"(?:95%\s*)?(?:bootstrap\s*)?CI\s*\[\s*([−\-+]?\d*\.?\d+)\s*,\s*([−\-+]?\d*\.?\d+)\s*\]", re.I),
    "test_statistics": re.compile(r"\b(F|t|z|χ²|chi-square|r|β|b|OR|RR|HR)\s*(?:\([^)]*\))?\s*=\s*([−\-+]?\d*\.?\d+)", re.I),
    "degrees_of_freedom": re.compile(r"\b(?:F|t|χ²)\s*\(([^)]*)\)", re.I),
    "percentages": re.compile(r"\b(\d{1,3}(?:\.\d+)?)%"),
    "effect_sizes": re.compile(r"\b(?:ηp²|η²|Cohen['’]?s?\s+d|d|f²|R²|ΔR²)\s*=\s*([−\-+]?\d*\.?\d+)", re.I),
}

STUDY_HEADING_RE = re.compile(r"^(?:study|experiment)\s+(\d+|[ivx]+)\b", re.I)


def all_blocks(doc: Document):
    for p in doc.paragraphs:
        if p.text.strip():
            yield p.text.strip()
    for table in doc.tables:
        for row in table.rows:
            text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip())
            if text:
                yield text


def norm_number(value: str) -> float:
    return float(value.replace("−", "-").replace(",", ""))


def audit(path: Path) -> dict:
    doc = Document(path)
    blocks = list(all_blocks(doc))
    current_study = "Front matter / unassigned"
    study_samples: dict[str, list[dict]] = defaultdict(list)
    inventory: dict[str, list[dict]] = {k: [] for k in PATTERNS}
    flags: list[dict] = []

    for idx, text in enumerate(blocks, start=1):
        heading = STUDY_HEADING_RE.match(text)
        if heading:
            current_study = f"Study {heading.group(1)}"
        for name, pattern in PATTERNS.items():
            for match in pattern.finditer(text):
                record = {"block": idx, "study": current_study, "match": match.group(0), "context": text[:360]}
                inventory[name].append(record)
                if name == "sample_sizes":
                    study_samples[current_study].append({"value": int(match.group(1).replace(",", "")), **record})
                elif name == "p_values":
                    op, raw = match.group(1), match.group(2)
                    value = float(raw if raw.startswith("0") else f"0{raw}" if raw.startswith(".") else raw)
                    if op == "=" and value == 0:
                        flags.append({"severity": "moderate", "type": "p-value", "block": idx, "issue": "Reported as p = 0; normally report p < .001 or exact software output.", "context": text[:360]})
                    if not 0 <= value <= 1:
                        flags.append({"severity": "serious", "type": "p-value", "block": idx, "issue": "p-value outside [0, 1].", "context": text[:360]})
                elif name == "confidence_intervals":
                    lo, hi = norm_number(match.group(1)), norm_number(match.group(2))
                    if lo > hi:
                        flags.append({"severity": "serious", "type": "confidence interval", "block": idx, "issue": "Lower confidence limit exceeds upper limit.", "context": text[:360]})
                elif name == "percentages":
                    value = float(match.group(1))
                    if value > 100:
                        flags.append({"severity": "serious", "type": "percentage", "block": idx, "issue": "Percentage exceeds 100%.", "context": text[:360]})

    sample_summary = {}
    for study, entries in study_samples.items():
        counts = Counter(e["value"] for e in entries)
        sample_summary[study] = {"values": sorted(counts), "occurrences": dict(sorted(counts.items()))}
        if len(counts) > 4:
            flags.append({
                "severity": "moderate",
                "type": "sample size variability",
                "study": study,
                "issue": "Many distinct sample-size values appear. Verify recruitment, exclusions, final N, analytic Ns, and cell sizes.",
                "values": sorted(counts),
            })

    missing_df_candidates = []
    for rec in inventory["test_statistics"]:
        if rec["match"].lower().startswith(("f", "t", "χ²", "chi-square")) and "(" not in rec["match"]:
            missing_df_candidates.append(rec)

    return {
        "file": str(path),
        "inventory_counts": {k: len(v) for k, v in inventory.items()},
        "sample_size_summary": sample_summary,
        "flags": flags,
        "possible_missing_degrees_of_freedom": missing_df_candidates[:100],
        "limitations": [
            "The script inventories reported values but cannot verify them against raw data or statistical output.",
            "Sample-size variation may be legitimate because recruited, final, cell, and model-specific Ns differ.",
            "Degrees-of-freedom detection is heuristic and may miss values split across runs or tables.",
        ],
    }


def markdown(data: dict) -> str:
    lines = ["# Statistical Reporting Audit", "", f"- File: `{data['file']}`", "", "## Inventory"]
    for k, v in data["inventory_counts"].items():
        lines.append(f"- {k.replace('_', ' ').title()}: {v}")
    lines += ["", "## Sample-Size Summary"]
    for study, info in data["sample_size_summary"].items():
        lines.append(f"- **{study}**: {info['values']}")
    if not data["sample_size_summary"]:
        lines.append("- No `N =` or `n =` patterns detected.")
    lines += ["", "## Flags"]
    for flag in data["flags"]:
        lines.append(f"- **{flag.get('severity', '').title()}**: {flag.get('issue')} {flag.get('context', '')}")
    if not data["flags"]:
        lines.append("- No automatic range or ordering errors detected.")
    lines += ["", "## Possible Missing Degrees of Freedom"]
    for rec in data["possible_missing_degrees_of_freedom"][:25]:
        lines.append(f"- Block {rec['block']}: {rec['context']}")
    if not data["possible_missing_degrees_of_freedom"]:
        lines.append("- None detected heuristically.")
    lines += ["", "## Limitations"] + [f"- {x}" for x in data["limitations"]]
    return "\n".join(lines) + "\n"


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--json", dest="json_path", type=Path)
    parser.add_argument("--markdown", dest="md_path", type=Path)
    args = parser.parse_args()
    data = audit(args.docx)
    raw = json.dumps(data, indent=2, ensure_ascii=False)
    if args.json_path:
        args.json_path.write_text(raw + "\n", encoding="utf-8")
    else:
        print(raw)
    if args.md_path:
        args.md_path.write_text(markdown(data), encoding="utf-8")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
