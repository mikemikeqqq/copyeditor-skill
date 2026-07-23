#!/usr/bin/env python3
"""Structural and text-integrity audit for academic DOCX manuscripts."""
from __future__ import annotations

import argparse
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

from lxml import etree
from docx import Document

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {"w": W, "r": R}

MOJIBAKE_PATTERNS = ["Ã", "Â", "â€", "â€™", "â€œ", "â€˜", "ï¿½", "�"]
CJK_RE = re.compile(r"[\u3400-\u4DBF\u4E00-\u9FFF\uF900-\uFAFF\u3040-\u30FF\uAC00-\uD7AF]")


def xml_root(zf: zipfile.ZipFile, name: str):
    try:
        return etree.fromstring(zf.read(name))
    except KeyError:
        return None


def visible_text_from_xml(root) -> str:
    if root is None:
        return ""
    texts = root.xpath(".//w:t/text() | .//w:delText/text()", namespaces=NS)
    return "".join(texts)


def count_xpath(root, expr: str) -> int:
    if root is None:
        return 0
    return len(root.xpath(expr, namespaces=NS))


def audit(path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"file": str(path), "exists": path.exists(), "valid_zip": False}
    if not path.exists():
        return result

    try:
        with zipfile.ZipFile(path) as zf:
            result["valid_zip"] = True
            names = set(zf.namelist())
            result["package_parts"] = len(names)
            doc_root = xml_root(zf, "word/document.xml")
            settings_root = xml_root(zf, "word/settings.xml")
            comments_root = xml_root(zf, "word/comments.xml")
            footnotes_root = xml_root(zf, "word/footnotes.xml")
            endnotes_root = xml_root(zf, "word/endnotes.xml")
            rels_root = xml_root(zf, "word/_rels/document.xml.rels")

            result["track_revisions_enabled"] = bool(
                settings_root is not None
                and settings_root.xpath(".//w:trackRevisions", namespaces=NS)
            )
            result["revisions"] = {
                "insertions": count_xpath(doc_root, ".//w:ins"),
                "deletions": count_xpath(doc_root, ".//w:del"),
                "moves_from": count_xpath(doc_root, ".//w:moveFrom"),
                "moves_to": count_xpath(doc_root, ".//w:moveTo"),
                "format_changes": sum(
                    count_xpath(doc_root, expr)
                    for expr in [".//w:rPrChange", ".//w:pPrChange", ".//w:tblPrChange", ".//w:tcPrChange"]
                ),
            }
            result["comments"] = {
                "count": count_xpath(comments_root, ".//w:comment"),
                "anchors_start": count_xpath(doc_root, ".//w:commentRangeStart"),
                "anchors_end": count_xpath(doc_root, ".//w:commentRangeEnd"),
                "references": count_xpath(doc_root, ".//w:commentReference"),
                "authors": Counter(
                    comments_root.xpath(".//w:comment/@w:author", namespaces=NS) if comments_root is not None else []
                ),
            }
            result["fields"] = {
                "field_chars": count_xpath(doc_root, ".//w:fldChar"),
                "simple_fields": count_xpath(doc_root, ".//w:fldSimple"),
                "instructions": count_xpath(doc_root, ".//w:instrText"),
            }
            result["structure"] = {
                "paragraphs_xml": count_xpath(doc_root, ".//w:p"),
                "tables_xml": count_xpath(doc_root, ".//w:tbl"),
                "sections": count_xpath(doc_root, ".//w:sectPr"),
                "drawings": count_xpath(doc_root, ".//w:drawing") + count_xpath(doc_root, ".//w:pict"),
                "footnotes": max(0, count_xpath(footnotes_root, ".//w:footnote") - 2),
                "endnotes": max(0, count_xpath(endnotes_root, ".//w:endnote") - 2),
                "headers": len([n for n in names if n.startswith("word/header") and n.endswith(".xml")]),
                "footers": len([n for n in names if n.startswith("word/footer") and n.endswith(".xml")]),
                "media_files": len([n for n in names if n.startswith("word/media/")]),
                "charts": len([n for n in names if n.startswith("word/charts/") and n.endswith(".xml")]),
                "embeddings": len([n for n in names if n.startswith("word/embeddings/")]),
            }
            hyperlink_rels = 0
            if rels_root is not None:
                hyperlink_rels = len(
                    rels_root.xpath(
                        ".//*[local-name()='Relationship'][contains(@Type, '/hyperlink')]"
                    )
                )
            result["structure"]["hyperlink_relationships"] = hyperlink_rels

            xml_text = visible_text_from_xml(doc_root)
            comment_text = visible_text_from_xml(comments_root)
            full_text = xml_text + "\n" + comment_text
            result["text_integrity"] = {
                "visible_characters": len(xml_text),
                "cjk_characters": len(CJK_RE.findall(full_text)),
                "mojibake_counts": {p: full_text.count(p) for p in MOJIBAKE_PATTERNS if p in full_text},
                "replacement_characters": full_text.count("�"),
            }

        doc = Document(path)
        result["python_docx"] = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "inline_shapes": len(doc.inline_shapes),
            "sections": len(doc.sections),
        }
        result["status"] = "ok"
    except Exception as exc:  # noqa: BLE001
        result["status"] = "error"
        result["error"] = f"{type(exc).__name__}: {exc}"
    return result


def to_markdown(data: dict[str, Any]) -> str:
    lines = ["# DOCX Audit", "", f"- File: `{data.get('file', '')}`", f"- Status: **{data.get('status', 'unknown')}**"]
    if data.get("error"):
        lines.append(f"- Error: `{data['error']}`")
        return "\n".join(lines) + "\n"
    lines += [
        f"- Valid ZIP: {data.get('valid_zip')}",
        f"- Track revisions enabled: {data.get('track_revisions_enabled')}",
        "",
        "## Revisions",
    ]
    for k, v in data.get("revisions", {}).items():
        lines.append(f"- {k.replace('_', ' ').title()}: {v}")
    lines += ["", "## Comments"]
    comments = data.get("comments", {})
    for k in ["count", "anchors_start", "anchors_end", "references"]:
        lines.append(f"- {k.replace('_', ' ').title()}: {comments.get(k, 0)}")
    lines.append(f"- Authors: {dict(comments.get('authors', {}))}")
    lines += ["", "## Fields"]
    for k, v in data.get("fields", {}).items():
        lines.append(f"- {k.replace('_', ' ').title()}: {v}")
    lines += ["", "## Structure"]
    for k, v in data.get("structure", {}).items():
        lines.append(f"- {k.replace('_', ' ').title()}: {v}")
    lines += ["", "## Text Integrity"]
    for k, v in data.get("text_integrity", {}).items():
        lines.append(f"- {k.replace('_', ' ').title()}: {v}")
    return "\n".join(lines) + "\n"


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--json", dest="json_path", type=Path)
    parser.add_argument("--markdown", dest="md_path", type=Path)
    args = parser.parse_args()

    data = audit(args.docx)
    output = json.dumps(data, indent=2, ensure_ascii=False, default=dict)
    if args.json_path:
        args.json_path.write_text(output + "\n", encoding="utf-8")
    else:
        print(output)
    if args.md_path:
        args.md_path.write_text(to_markdown(data), encoding="utf-8")
    return 0 if data.get("status") == "ok" else 1


if __name__ == "__main__":
    raise SystemExit(main())
