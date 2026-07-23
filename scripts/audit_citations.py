#!/usr/bin/env python3
"""Heuristic author-year citation/reference correspondence audit for DOCX manuscripts."""
from __future__ import annotations

import argparse
import json
import re
from collections import Counter, defaultdict
from pathlib import Path
from typing import Iterable

from docx import Document

YEAR_RE = r"(?:19|20)\d{2}[a-z]?"
PAREN_CLUSTER_RE = re.compile(r"\(([^()]*?(?:19|20)\d{2}[a-z]?[^()]*)\)")
NARRATIVE_RE = re.compile(
    rf"(?P<author>[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿ'’\-]+(?:\s+(?:et\s+al\.|(?:&|and)\s+[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿ'’\-]+(?:['’]s)?))?)\s*\((?P<year>{YEAR_RE})\)"
)
REF_START_RE = re.compile(rf"^(?P<author>[A-ZÀ-ÖØ-Þ][^\n]{{0,180}}?)\s*\((?P<year>{YEAR_RE})\)\.")


def iter_text(doc: Document) -> Iterable[str]:
    for p in doc.paragraphs:
        if p.text.strip():
            yield p.text.strip()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if text:
                    yield text


def first_author(author_text: str) -> str:
    cleaned = re.sub(r"\bet\s+al\.\b", "", author_text, flags=re.I)
    cleaned = cleaned.replace("&", ",")
    token = re.split(r"[,\s]", cleaned.strip())[0]
    return re.sub(r"[^A-Za-zÀ-ÖØ-öø-ÿ'’\-]", "", token).lower()


def key(author_text: str, year: str) -> str:
    return f"{first_author(author_text)}|{year.lower()}"


def find_reference_index(paragraphs: list[str]) -> int | None:
    candidates = []
    for i, text in enumerate(paragraphs):
        norm = re.sub(r"\s+", " ", text.strip().lower()).rstrip(":")
        if norm in {"references", "reference list", "bibliography"}:
            candidates.append(i)
    return candidates[0] if candidates else None


def audit(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    ref_idx = find_reference_index(paragraphs)
    body = paragraphs[:ref_idx] if ref_idx is not None else paragraphs
    refs = paragraphs[ref_idx + 1 :] if ref_idx is not None else []

    in_text = []
    for text in body:
        for cluster in PAREN_CLUSTER_RE.findall(text):
            for segment in cluster.split(";"):
                years = list(re.finditer(YEAR_RE, segment))
                if not years:
                    continue
                for ym in years:
                    prefix = segment[: ym.start()]
                    candidates = re.findall(r"[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿ'’\-]+", prefix)
                    if not candidates:
                        continue
                    author = candidates[0]
                    year = ym.group(0)
                    in_text.append({"author": author, "year": year, "key": key(author, year), "context": text[:260]})
        for m in NARRATIVE_RE.finditer(text):
            in_text.append({"author": m.group("author"), "year": m.group("year"), "key": key(m.group("author"), m.group("year")), "context": text[:260]})

    reference_entries = []
    for text in refs:
        m = REF_START_RE.match(text)
        if m:
            reference_entries.append({"author": m.group("author"), "year": m.group("year"), "key": key(m.group("author"), m.group("year")), "entry": text})

    cited_keys = Counter(x["key"] for x in in_text)
    ref_keys = Counter(x["key"] for x in reference_entries)
    missing_refs = sorted(k for k in cited_keys if k and k not in ref_keys)
    uncited_refs = sorted(k for k in ref_keys if k and k not in cited_keys)
    duplicate_refs = sorted(k for k, n in ref_keys.items() if n > 1)

    contexts = defaultdict(list)
    for item in in_text:
        contexts[item["key"]].append(item["context"])

    malformed_dois = []
    for entry in reference_entries:
        text = entry["entry"]
        if "doi" in text.lower():
            doi_tokens = re.findall(r"(?:https?://(?:dx\.)?doi\.org/|doi:\s*)(\S+)", text, flags=re.I)
            for token in doi_tokens:
                if not re.match(r"10\.\d{4,9}/\S+", token.rstrip(".,;")):
                    malformed_dois.append(text)

    return {
        "file": str(path),
        "reference_heading_found": ref_idx is not None,
        "in_text_citation_occurrences": len(in_text),
        "unique_in_text_keys": len(cited_keys),
        "reference_entries_parsed": len(reference_entries),
        "unique_reference_keys": len(ref_keys),
        "missing_reference_keys": missing_refs,
        "missing_reference_examples": {k: contexts[k][:3] for k in missing_refs},
        "uncited_reference_keys": uncited_refs,
        "duplicate_reference_keys": duplicate_refs,
        "malformed_doi_entries": malformed_dois,
        "limitations": [
            "Author-year matching is heuristic and uses first-author surname plus year.",
            "Institutional authors, legal cases, suffixes, and citation fields require manual verification.",
            "Appendix-only reference lists may be classified as uncited from the main text.",
        ],
    }


def markdown(data: dict) -> str:
    lines = [
        "# Citation and Reference Audit",
        "",
        f"- File: `{data['file']}`",
        f"- Reference heading found: {data['reference_heading_found']}",
        f"- In-text citation occurrences: {data['in_text_citation_occurrences']}",
        f"- Parsed reference entries: {data['reference_entries_parsed']}",
        "",
        "## Potential Missing References",
    ]
    lines += [f"- `{x}`" for x in data["missing_reference_keys"]] or ["- None detected heuristically."]
    lines += ["", "## Potential Uncited References"]
    lines += [f"- `{x}`" for x in data["uncited_reference_keys"]] or ["- None detected heuristically."]
    lines += ["", "## Duplicate Author-Year Keys"]
    lines += [f"- `{x}`" for x in data["duplicate_reference_keys"]] or ["- None detected heuristically."]
    lines += ["", "## DOI Flags"]
    lines += [f"- {x}" for x in data["malformed_doi_entries"]] or ["- None detected heuristically."]
    lines += ["", "## Limitations"] + [f"- {x}" for x in data["limitations"]]
    return "\n".join(lines) + "\n"


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--json", dest="json_path", type=Path)
    parser.add_argument("--markdown", dest="md_path", type=Path)
    args = parser.parse_args()
    data = audit(args.docx)
    raw = json.dumps(data, indent=2, ensure_ascii=False)
    if args.json_path:
        args.json_path.write_text(raw + "\n", encoding="utf-8")
    else:
        print(raw)
    if args.md_path:
        args.md_path.write_text(markdown(data), encoding="utf-8")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
